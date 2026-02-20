"""
Document Processor Module for Ask-Chopper

Handles document text extraction, chunking, and embedding generation.
Supports PDF, DOCX, TXT, and various code/text file formats.
"""

import os
import uuid
import io
from typing import Tuple, List, Optional
from sentence_transformers import SentenceTransformer

# Embedding model configuration (local model, no external API key required)
EMBEDDING_MODEL = os.environ.get("EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
EMBEDDING_DIMENSIONS = 384

_embedding_model = None

def _get_embedding_model():
    """Get or create local embedding model."""
    global _embedding_model
    if _embedding_model is None:
        _embedding_model = SentenceTransformer(EMBEDDING_MODEL)
    return _embedding_model

# Chunking configuration
DEFAULT_CHUNK_SIZE = 1000  # characters - larger chunks for better context
DEFAULT_OVERLAP = 100  # characters - more overlap for continuity


def extract_text(file, mime_type: str) -> str:
    """
    Extract text content from uploaded file.

    Args:
        file: File object (werkzeug FileStorage or file-like object)
        mime_type: MIME type of the file

    Returns:
        Extracted text content
    """
    # Read file content - handle different file object types
    try:
        # Try to reset file pointer first
        if hasattr(file, 'seek'):
            file.seek(0)

        # Read content
        if hasattr(file, 'read'):
            content = file.read()
        elif hasattr(file, 'stream'):
            # Werkzeug FileStorage
            file.stream.seek(0)
            content = file.stream.read()
        else:
            raise ValueError("Unable to read file content")

        # Ensure we have bytes
        if isinstance(content, str):
            content = content.encode('utf-8')

        print(f"DEBUG extract_text: Read {len(content)} bytes from file")

    except Exception as e:
        print(f"ERROR reading file: {e}")
        return f"[Error reading file: {str(e)}]"

    # PDF extraction
    if mime_type == "application/pdf" or (hasattr(file, 'filename') and file.filename.lower().endswith('.pdf')):
        return _extract_pdf(content)

    # DOCX extraction
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or \
       (hasattr(file, 'filename') and file.filename.lower().endswith('.docx')):
        return _extract_docx(content)

    # DOC (older format) - try as text or show error
    if mime_type == "application/msword" or \
       (hasattr(file, 'filename') and file.filename.lower().endswith('.doc')):
        # .doc format requires python-docx-binary or antiword - fallback to error message
        return "[Error: Legacy .doc format not fully supported. Please convert to .docx or .pdf]"

    # Plain text and code files
    text_mimes = [
        "text/plain", "text/markdown", "text/html", "text/css",
        "text/csv", "text/xml", "application/json", "application/xml",
        "application/javascript", "text/javascript"
    ]
    text_extensions = ['.txt', '.md', '.py', '.js', '.json', '.csv', '.xml', '.html', '.css', '.yml', '.yaml']

    if mime_type in text_mimes or \
       (hasattr(file, 'filename') and any(file.filename.lower().endswith(ext) for ext in text_extensions)):
        return _extract_text(content)

    # Fallback: try to decode as text
    return _extract_text(content)


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF content using pdfplumber (primary) or PyPDF2 (fallback)."""
    print(f"DEBUG _extract_pdf: Processing PDF with {len(content)} bytes")

    # Try pdfplumber first - better for complex PDFs
    try:
        import pdfplumber

        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            print(f"DEBUG _extract_pdf: PDF has {len(pdf.pages)} pages (pdfplumber)")

            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                        print(f"DEBUG _extract_pdf: Page {page_num + 1} extracted {len(page_text)} chars")
                    else:
                        # Try extracting tables if no text
                        tables = page.extract_tables()
                        if tables:
                            table_text = []
                            for table in tables:
                                for row in table:
                                    row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                    table_text.append(row_text)
                            if table_text:
                                text_parts.append(f"[Page {page_num + 1} - Table]\n" + "\n".join(table_text))
                                print(f"DEBUG _extract_pdf: Page {page_num + 1} extracted table data")
                        else:
                            print(f"DEBUG _extract_pdf: Page {page_num + 1} has no extractable text")
                except Exception as page_error:
                    print(f"ERROR extracting page {page_num + 1}: {page_error}")

        result = "\n\n".join(text_parts)
        print(f"DEBUG _extract_pdf: Total extracted text: {len(result)} chars (pdfplumber)")

        if result:
            return result

    except Exception as e:
        print(f"DEBUG _extract_pdf: pdfplumber failed: {e}, trying PyPDF2...")

    # Fallback to PyPDF2
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(content))
        print(f"DEBUG _extract_pdf: PDF has {len(reader.pages)} pages (PyPDF2)")

        text_parts = []

        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                    print(f"DEBUG _extract_pdf: Page {page_num + 1} extracted {len(page_text)} chars")
            except Exception as page_error:
                print(f"ERROR extracting page {page_num + 1}: {page_error}")

        result = "\n\n".join(text_parts)
        print(f"DEBUG _extract_pdf: Total extracted text: {len(result)} chars (PyPDF2)")

        if result:
            return result

    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"[Error extracting PDF text: {str(e)}]"

    return "[No extractable text found in PDF. The PDF may contain only images or scanned content.]"


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX content."""
    try:
        from docx import Document

        print(f"DEBUG _extract_docx: Processing DOCX with {len(content)} bytes")
        doc = Document(io.BytesIO(content))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_rows.append(row_text)
            if table_rows:
                text_parts.append("[Table]\n" + "\n".join(table_rows))

        result = "\n\n".join(text_parts)
        print(f"DEBUG _extract_docx: Extracted {len(result)} chars from DOCX")

        if not result:
            return "[No text content found in DOCX document.]"

        return result
    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"[Error extracting DOCX text: {str(e)}]"


def _extract_text(content: bytes) -> str:
    """Extract text from plain text content."""
    print(f"DEBUG _extract_text: Processing text with {len(content)} bytes")

    # Try common encodings
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

    for encoding in encodings:
        try:
            result = content.decode(encoding)
            print(f"DEBUG _extract_text: Decoded with {encoding}, got {len(result)} chars")
            return result
        except (UnicodeDecodeError, LookupError):
            continue

    # Fallback: decode with errors ignored
    result = content.decode('utf-8', errors='ignore')
    print(f"DEBUG _extract_text: Fallback decode, got {len(result)} chars")
    return result


def chunk_text(
    text: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_OVERLAP
) -> List[str]:
    """
    Split text into overlapping chunks.

    Args:
        text: Full text to chunk
        chunk_size: Maximum size of each chunk in characters
        overlap: Number of overlapping characters between chunks

    Returns:
        List of text chunks
    """
    if not text or len(text) <= chunk_size:
        return [text] if text else []

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # If not at the end, try to break at a sentence or word boundary
        if end < len(text):
            # Look for sentence boundary (. ! ? followed by space or newline)
            for boundary in ['. ', '.\n', '! ', '!\n', '? ', '?\n', '\n\n']:
                last_boundary = text[start:end].rfind(boundary)
                if last_boundary > chunk_size * 0.5:  # Only break if past halfway
                    end = start + last_boundary + len(boundary)
                    break
            else:
                # No sentence boundary found, try word boundary
                last_space = text[start:end].rfind(' ')
                if last_space > chunk_size * 0.5:
                    end = start + last_space + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        # Move start position with overlap
        start = end - overlap if end < len(text) else end

    return chunks


def generate_embeddings(texts: List[str]) -> List[List[float]]:
    """
    Generate embeddings for a list of texts using a local embedding model.

    Args:
        texts: List of text strings to embed

    Returns:
        List of embedding vectors
    """
    if not texts:
        return []

    embedding_model = _get_embedding_model()
    vectors = embedding_model.encode(texts, normalize_embeddings=True).tolist()
    return vectors


def generate_query_embedding(query: str) -> List[float]:
    """
    Generate embedding for a single query string.

    Args:
        query: Query text

    Returns:
        Embedding vector
    """
    embedding_model = _get_embedding_model()
    vector = embedding_model.encode([query], normalize_embeddings=True).tolist()
    return vector[0]


def process_document(
    file,
    user_id: int,
    session_id: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_OVERLAP
) -> Tuple[str, List[str], List[List[float]]]:
    """
    Full document processing pipeline.

    Args:
        file: File object to process
        user_id: User ID for metadata
        session_id: Session ID for metadata
        chunk_size: Size of text chunks
        overlap: Overlap between chunks

    Returns:
        Tuple of (doc_id, chunks, embeddings)
    """
    # Generate unique document ID
    doc_id = str(uuid.uuid4())

    # Get MIME type
    mime_type = getattr(file, 'content_type', None)
    if not mime_type:
        import mimetypes
        mime_type = mimetypes.guess_type(file.filename)[0] or 'application/octet-stream'

    # Extract text
    text = extract_text(file, mime_type)

    if not text or text.startswith('[Error'):
        raise ValueError(f"Failed to extract text from document: {text}")

    # Chunk text
    chunks = chunk_text(text, chunk_size, overlap)

    if not chunks:
        raise ValueError("No text chunks generated from document")

    # Generate embeddings
    embeddings = generate_embeddings(chunks)

    return doc_id, chunks, embeddings


def estimate_tokens(text: str) -> int:
    """
    Estimate token count for a text string.
    Uses tiktoken for accurate counting if available.

    Args:
        text: Text to count tokens for

    Returns:
        Estimated token count
    """
    try:
        import tiktoken
        encoding = tiktoken.get_encoding("cl100k_base")
        return len(encoding.encode(text))
    except Exception:
        # Fallback: rough estimate (1 token ~ 4 characters)
        return len(text) // 4


def build_context_prompt(
    query: str,
    retrieved_chunks: List[str],
    max_context_tokens: int = 4000
) -> str:
    """
    Build a context-augmented prompt from retrieved chunks.

    Args:
        query: User's query
        retrieved_chunks: List of relevant document chunks
        max_context_tokens: Maximum tokens for context section

    Returns:
        Formatted prompt with context
    """
    if not retrieved_chunks:
        return query

    # Build context section with token budget
    context_parts = []
    total_tokens = 0

    for i, chunk in enumerate(retrieved_chunks):
        chunk_tokens = estimate_tokens(chunk)
        if total_tokens + chunk_tokens > max_context_tokens:
            break
        context_parts.append(f"[Section {i + 1}]:\n{chunk}")
        total_tokens += chunk_tokens

    if not context_parts:
        return query

    context = "\n\n".join(context_parts)

    prompt = f"""I have uploaded a document. Here are the relevant sections from it:

--- DOCUMENT CONTENT ---
{context}
--- END OF DOCUMENT ---

My question/request: {query}

Please analyze the document content above and respond to my request. If I'm asking you to explain or summarize, provide a clear and comprehensive explanation. If I'm asking a specific question, answer based on the document content."""

    return prompt
